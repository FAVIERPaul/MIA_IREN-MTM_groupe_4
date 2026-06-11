"""Phase 5 — Génération automatisée de lettre de motivation avec Cerebras.

Diffère du notebook original : ici PAS d'aller-retour manuel sur claude.ai. La LM est générée
directement par gpt-oss-120b avec le mégaprompt + few-shot des LM exemples de l'utilisateur.
"""
from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path
from typing import Callable

from docx import Document
from docx.shared import Cm, Pt

from . import config
from .cerebras_client import call_llm


# ─── Étape 1 : fiche entreprise ─────────────────────────────────────
_SYSTEM_FICHE = """Tu es analyste en intelligence économique. Produis une fiche synthétique sur une
entreprise, pour aider à rédiger une lettre de motivation PERSONNALISÉE.

Identifie :
- Activité principale, positionnement, secteur
- Valeurs et culture (telles qu'elles ressortent du discours)
- 1-3 projets/initiatives récents NOMINATIVEMENT IDENTIFIABLES (rachats, lancements,
  partenariats, déclarations dirigeants)
- Vocabulaire distinctif (mots qui reviennent dans leur com)
- Ton de communication
- Une ACCROCHE PERSONNALISEE précise utilisable dans le paragraphe 2 de la LM
  (= un détail nominatif que seul un candidat qui s'est renseigné pourrait citer)

Réponds UNIQUEMENT en JSON :
{
  "activite": "",
  "positionnement": "",
  "valeurs": [],
  "projets_recents": [{"projet": "", "details": ""}],
  "vocabulaire_distinctif": [],
  "ton_communication": "",
  "accroche_personnalisee_suggeree": "Une phrase prête à insérer dans la LM"
}"""


def build_fiche_entreprise(offre: dict, website_text: str = "", articles_text: str = "") -> dict:
    user = (
        f"ENTREPRISE : {offre.get('entreprise', '')}\n"
        f"OFFRE     : {offre.get('intitule', '')}\n\n"
        f"DESCRIPTION DE L'OFFRE :\n{(offre.get('description') or '')[:2000]}\n\n"
        f"SITE WEB (extrait) :\n{website_text[:3000]}\n\n"
        f"ARTICLES :\n{articles_text[:3000]}"
    )
    return call_llm(
        messages=[
            {"role": "system", "content": _SYSTEM_FICHE},
            {"role": "user",   "content": user},
        ],
        max_tokens=2500,
    )


# ─── Étape 2 : audit angles stratégiques ─────────────────────────────
_SYSTEM_ANGLES = """Tu es expert en stratégie de candidature.

À partir du profil du candidat (CV détaillé + profil enrichi) et de l'offre + fiche
entreprise, identifie LES ANGLES STRATEGIQUES pour rédiger une LM percutante :

1. EXPERIENCE PIVOT : LA seule expérience à mettre massivement en avant.
   Critère : max match entre les responsabilités du poste et les missions réelles
   du candidat.

2. TON A ADOPTER (1 seul choix) :
   - 'chaleureux_passionne'   → culture / média / ONG / audiovisuel
   - 'formel_intellectuel'    → conseil / finance / luxe / recherche
   - 'dynamique_operationnel' → startup / tech / événementiel / com / marketing
   - 'neutre_corporate'       → grand groupe B2B / industrie

3. REFERENCE NOMINATIVE : un projet/citation/initiative PRECISE issue de la fiche
   entreprise. C'est LE détail qui prouve que le candidat s'est vraiment renseigné.

4. 2-3 MOTS-CLES ATS CRITIQUES à intégrer naturellement.

5. TWIST PERSONNEL : un élément du profil qui humanise. Choisis celui qui résonne
   le plus avec le secteur.

Réponds UNIQUEMENT en JSON :
{
  "experience_pivot": {"intitule": "", "entreprise": "", "pourquoi": ""},
  "ton": "chaleureux_passionne | formel_intellectuel | dynamique_operationnel | neutre_corporate",
  "ton_justification": "",
  "reference_nominative": "Phrase concrète à utiliser",
  "mots_cles_ats": ["", "", ""],
  "twist_personnel": ""
}"""


def identify_angles(offre: dict, fiche: dict, profil: dict, cv_text_brut: str) -> dict:
    user = (
        f"OFFRE :\nIntitulé   : {offre.get('intitule', '')}\n"
        f"Entreprise : {offre.get('entreprise', '')}\n"
        f"Description :\n{(offre.get('description') or '')[:2500]}\n\n"
        f"FICHE ENTREPRISE :\n{json.dumps(fiche, ensure_ascii=False, indent=2)}\n\n"
        f"PROFIL ENRICHI :\n{json.dumps(profil, ensure_ascii=False, indent=2)}\n\n"
        f"CV BRUT (pour identifier l'expérience pivot) :\n{cv_text_brut[:3500]}"
    )
    return call_llm(
        messages=[
            {"role": "system", "content": _SYSTEM_ANGLES},
            {"role": "user",   "content": user},
        ],
        max_tokens=2500,
    )


# ─── Étape 3 : génération de la LM (TEXTE) ──────────────────────────
_TON_DIRECTIVES = {
    "chaleureux_passionne": (
        "Ton CHALEUREUX et PERSONNEL. Tu peux montrer une connexion sincère au secteur "
        "(culture, média). Phrases plus douces, un peu plus longues. Le 'je' peut être plus "
        "assumé émotionnellement. Référence à une expérience personnelle de spectateur/auditeur OK."
    ),
    "formel_intellectuel": (
        "Ton FORMEL et INTELLECTUEL. Vocabulaire précis, presque académique. Exemples : "
        "'défi intellectuel', 'grille de lecture', 'point de vue décentré'. Phrases construites, "
        "syntaxe maîtrisée. Une touche humanisante via une discipline exigeante."
    ),
    "dynamique_operationnel": (
        "Ton DYNAMIQUE et OPERATIONNEL. Phrases plus courtes, plus directes. Vocabulaire "
        "action : 'piloter', 'coordonner', 'force de proposition', 'immédiatement opérationnel'. "
        "Mise en avant de résultats concrets."
    ),
    "neutre_corporate": (
        "Ton NEUTRE et CORPORATE. Pas trop chaleureux, pas trop intello. Vocabulaire business "
        "standard. Mise en avant de la rigueur, de la fiabilité, de l'adaptabilité."
    ),
}


def _build_few_shot_block(lm_files: list[Path]) -> str:
    if not lm_files:
        return "(Aucun exemple fourni — invente un style sobre et professionnel.)"
    blocks = []
    for i, p in enumerate(lm_files, 1):
        try:
            doc = Document(str(p))
            text = "\n".join(par.text for par in doc.paragraphs if par.text.strip())
            blocks.append(f"### EXEMPLE {i} — {p.stem}\n\n{text}")
        except Exception:
            continue
    return "\n\n---\n\n".join(blocks)


def generate_lm_text(
    offre: dict,
    profil: dict,
    fiche: dict,
    angles: dict,
    candidat: dict,
    cv_text_brut: str,
    lm_examples: list[Path] | None = None,
) -> str:
    """Génère la lettre de motivation en TEXTE BRUT via Cerebras."""
    lm_examples = lm_examples or sorted(config.LM_EXAMPLES_DIR.glob("*.docx"))
    fewshot = _build_few_shot_block(lm_examples)

    ton_choisi = angles.get("ton", "neutre_corporate")
    directives_ton = _TON_DIRECTIVES.get(ton_choisi, _TON_DIRECTIVES["neutre_corporate"])

    signature = (
        f"{candidat.get('nom_complet', '')}\n"
        f"{candidat.get('adresse', '')}\n"
        f"{candidat.get('telephone', '')}\n"
        f"{candidat.get('email', '')}"
    )

    system = (
        f"Tu es un expert en rédaction de lettres de motivation en français.\n\n"
        f"Tu rédiges une LM signée {candidat.get('nom_complet', '')} pour le poste donné. "
        f"Tu reproduis FIDÈLEMENT le style des {len(lm_examples)} LM exemples fournies (si présentes), "
        f"tout en adaptant le ton au secteur.\n\n"
        f"DIRECTIVES DE TON :\n{directives_ton}\n\n"
        "STRUCTURE :\n"
        f"1. EN-TÊTE : nom complet, adresse, téléphone, email (chacun sur sa propre ligne)\n"
        "2. Objet : 'Candidature au poste de [intitulé exact de l'offre]'\n"
        "3. Civilité : 'Madame, Monsieur,'\n"
        "4. Paragraphe 1 (3-4 lignes) : amorce + accroche personnalisée NOMINATIVE sur l'entreprise\n"
        "5. Paragraphe 2 (5-7 lignes) : expérience pivot avec détails concrets (chiffres, noms, contexte)\n"
        "6. Paragraphe 3 (3-5 lignes) : compétences complémentaires + twist personnel humanisant\n"
        "7. Paragraphe 4 (2-3 lignes) : ouverture sur un entretien + formule de politesse\n"
        f"8. Signature : {candidat.get('nom_complet', '')}\n\n"
        "RÈGLES :\n"
        "- 380-450 mots au total\n"
        "- N'invente jamais une expérience absente du CV\n"
        "- Intègre les mots-clés ATS naturellement, pas en liste\n"
        "- Cite la référence nominative donnée par la stratégie\n"
        "- Pas de markdown, pas de **bold**, juste du texte brut formaté\n"
        "- Rends UNIQUEMENT le texte de la LM, rien d'autre"
    )

    user = (
        f"=== OFFRE ===\nIntitulé : {offre.get('intitule', '')}\n"
        f"Entreprise : {offre.get('entreprise', '')}\n"
        f"Description :\n{(offre.get('description') or '')[:2000]}\n\n"
        f"=== FICHE ENTREPRISE ===\n{json.dumps(fiche, ensure_ascii=False, indent=2)}\n\n"
        f"=== STRATEGIE ===\n{json.dumps(angles, ensure_ascii=False, indent=2)}\n\n"
        f"=== PROFIL ENRICHI ===\n{json.dumps(profil, ensure_ascii=False, indent=2)}\n\n"
        f"=== CV BRUT (vérité ultime, ne rien inventer en-dehors) ===\n{cv_text_brut[:3500]}\n\n"
        f"=== EXEMPLES DE STYLE (à imiter) ===\n{fewshot}\n\n"
        f"=== EN-TÊTE À METTRE EN HAUT DE LA LM ===\n{signature}\n\n"
        "Génère la lettre maintenant."
    )

    # ⚙️  FIX 8 juin 2026 :
    #   reasoning_effort="medium" + max_tokens=2500 → finish_reason=length systématique
    #   (le reasoning interne mangeait tout le budget avant l'output).
    #   Combo "low" + 4000 tokens = LM écrite proprement, pas de reasoning superflu.
    return call_llm(
        messages=[
            {"role": "system", "content": system},
            {"role": "user",   "content": user},
        ],
        max_tokens=4000,
        temperature=0.5,
        json_mode=False,
        reasoning_effort="low",
    )


# ─── Étape 4 : critique automatique (optionnelle) ────────────────────
_SYSTEM_CRITIQUE = """Tu critiques une LM sur 5 axes objectifs.

1. REFERENCE NOMINATIVE : la LM cite-t-elle un projet/initiative PRECISE ?
2. EXPERIENCE PIVOT : expérience phare mise en avant avec détails concrets ?
3. MOTS-CLES ATS : intégrés naturellement ?
4. LONGUEUR : idéalement 380-450 mots.
5. TON : adapté au secteur ? Cohérent ?

Réponds UNIQUEMENT en JSON :
{
  "score_global": 0,
  "axes": {
    "reference_nominative": {"score": 0, "commentaire": ""},
    "experience_pivot":     {"score": 0, "commentaire": ""},
    "mots_cles_ats":        {"score": 0, "mots_presents": [], "mots_absents": []},
    "longueur":             {"nb_mots": 0, "verdict": "OK | trop court | trop long"},
    "ton":                  {"score": 0, "commentaire": ""}
  },
  "suggestions_amelioration": [],
  "verdict_final": "GO | A REVOIR"
}
(scores sur 10)"""


def critique_lm(lm_text: str, offre: dict, angles: dict) -> dict:
    user = (
        f"OFFRE : {offre.get('intitule', '')} chez {offre.get('entreprise', '')}\n\n"
        f"MOTS-CLES ATS ATTENDUS : {angles.get('mots_cles_ats', [])}\n"
        f"TON ATTENDU : {angles.get('ton', '')}\n\n"
        f"LM A CRITIQUER :\n\n{lm_text}"
    )
    return call_llm(
        messages=[
            {"role": "system", "content": _SYSTEM_CRITIQUE},
            {"role": "user",   "content": user},
        ],
        max_tokens=2500,
    )


# ─── Étape 5 : export .docx ──────────────────────────────────────────
def export_to_docx(lm_text: str, output_path: Path, template_source: Path | None = None) -> Path:
    """Exporte la LM en .docx, en clonant le format d'une LM exemple si dispo."""
    # Garde-fou supplémentaire : si jamais on arrive ici avec un truc invalide,
    # on lève une erreur claire plutôt qu'un AttributeError obscur.
    if not lm_text or not isinstance(lm_text, str) or not lm_text.strip():
        raise RuntimeError(
            "Impossible d'exporter une LM vide. "
            "La génération a probablement échoué en amont — vérifie l'étape 3."
        )

    if template_source and template_source.exists():
        doc = Document(str(template_source))
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)
    else:
        doc = Document()
        for s in doc.sections:
            s.top_margin    = Cm(1.27)
            s.bottom_margin = Cm(1.27)
            s.left_margin   = Cm(1.27)
            s.right_margin  = Cm(1.27)

    def add_par(text: str):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.15
        pf.space_after = Pt(0)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    # Découpe en blocs séparés par lignes vides
    lines = lm_text.strip().split("\n")
    current: list[str] = []
    blocks: list[str] = []
    for line in lines:
        if line.strip() == "":
            if current:
                blocks.append("\n".join(current))
                current = []
        else:
            current.append(line.rstrip())
    if current:
        blocks.append("\n".join(current))

    for block in blocks:
        if "\n" in block and len(block.split("\n")) <= 5:
            for line in block.split("\n"):
                add_par(line)
            add_par("")
        else:
            add_par(block)
            add_par("")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ─── Pipeline complet ───────────────────────────────────────────────
def generate_lm_for_offre(
    offre: dict,
    profil: dict,
    candidat: dict,
    cv_pdf_path: Path = config.CV_PDF_PATH,
    website_text: str = "",
    articles_text: str = "",
    progress_cb: Callable[[str], None] | None = None,
) -> dict:
    """Pipeline LM complet.
    Retourne {fiche, angles, lm_text, critique, docx_path}.
    """
    from .cv_generator import parse_cv_pdf, slugify

    def step(msg):
        if progress_cb:
            progress_cb(msg)

    step("Construction de la fiche entreprise…")
    fiche = build_fiche_entreprise(offre, website_text, articles_text)

    step("Identification des angles stratégiques…")
    cv_text = parse_cv_pdf(cv_pdf_path)
    angles = identify_angles(offre, fiche, profil, cv_text)

    step("Rédaction de la lettre…")
    lm_text = generate_lm_text(offre, profil, fiche, angles, candidat, cv_text)

    step("Auto-critique…")
    try:
        critique = critique_lm(lm_text, offre, angles)
    except Exception:
        critique = None

    step("Export .docx…")
    folder = (
        datetime.now().strftime("%Y%m%d_%H%M%S") + "_"
        + slugify(offre.get("entreprise", "")) + "_"
        + slugify(offre.get("intitule", ""))
    )
    out_dir = config.OUTPUTS_LM_DIR / folder
    docx_path = out_dir / "lettre_motivation.docx"

    # Clone le format d'une LM exemple si dispo (Times New Roman, marges 1.27cm)
    lm_files = sorted(config.LM_EXAMPLES_DIR.glob("*.docx"))
    template = lm_files[0] if lm_files else None
    export_to_docx(lm_text, docx_path, template)

    # Sauve aussi le texte brut pour réédition rapide
    (out_dir / "lettre_motivation.txt").write_text(lm_text, encoding="utf-8")

    return {
        "fiche":     fiche,
        "angles":    angles,
        "lm_text":   lm_text,
        "critique":  critique,
        "docx_path": docx_path,
        "out_dir":   out_dir,
    }
